from flask import Flask, render_template, request, redirect, url_for, session, flash
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import date, datetime
from functools import wraps
from io import BytesIO
from flask import send_file
from docx import Document
from docx.shared import Pt


app = Flask(__name__)

# 🔐 Cambialo con una stringa lunga e casuale
app.config['SECRET_KEY'] = 'cambia-questa-chiave-segreta'

# 🔧 Metti qui i dati del tuo MySQL su Aruba
# formato: mysql+pymysql://utente:password@host/nome_database
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://Sql1904907:Pierino_68@31.11.38.14/Sql1904907_1'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ==================
# MODELLI
# ==================

class Utente(db.Model):
    __tablename__ = 'utenti'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    ruolo = db.Column(db.String(50), default='segreteria')  # 'segreteria' | 'sostituzioni' | 'admin'

    def verifica_password(self, password):
        return check_password_hash(self.password_hash, password)


class Docente(db.Model):
    __tablename__ = 'docenti'
    id = db.Column(db.Integer, primary_key=True)
    cognome = db.Column(db.String(100), nullable=False)
    nome = db.Column(db.String(100))
    codice = db.Column(db.String(50), unique=True)
    attivo = db.Column(db.Boolean, default=True)
    fittizio = db.Column(db.Boolean, default=False)


class Assenza(db.Model):
    __tablename__ = 'assenze'
    id = db.Column(db.Integer, primary_key=True)
    docente_id = db.Column(db.Integer, db.ForeignKey('docenti.id'), nullable=False)
    data = db.Column(db.Date, nullable=False)
    ore = db.Column(db.String(50), nullable=False)  # es: "1,2,3"
    note = db.Column(db.Text)

    docente = db.relationship('Docente')


class Lezione(db.Model):
    __tablename__ = 'lezioni'
    id = db.Column(db.Integer, primary_key=True)
    docente_id = db.Column(db.Integer, db.ForeignKey('docenti.id'), nullable=False)
    activity_id = db.Column(db.Integer)
    day = db.Column(db.String(10), nullable=False)       # es: LUN, MAR, ...
    hour = db.Column(db.String(10), nullable=False)      # es: H1, H2, ...
    students_set = db.Column(db.String(50))              # classe / gruppo
    subject = db.Column(db.String(100))                  # materia
    room = db.Column(db.String(50))                      # aula
    note = db.Column(db.String(255))                     # commenti

    docente = db.relationship('Docente')


class Sostituzione(db.Model):
    __tablename__ = "sostituzioni"

    id = db.Column(db.Integer, primary_key=True)
    data = db.Column(db.Date, nullable=False)
    hour = db.Column(db.String(3), nullable=False)
    classe = db.Column(db.String(20))
    docente_assente_id = db.Column(db.Integer)
    docente_sostituto_id = db.Column(db.Integer)
    tipo = db.Column(db.String(1), default="D")

    entrata_posticipata = db.Column(db.Boolean, default=False)
    uscita_anticipata = db.Column(db.Boolean, default=False)

    created_at = db.Column(db.DateTime, server_default=db.func.now())

# ==================
# UTILITY
# ==================
def candidati_per_scopertura(giorno_data, giorno_cod, hour_tag, classe, docente_assente_id, assenti_ids):
    """
    Ritorna lista candidati come:
      [{"docente": <Docente>, "tipo": "C|D|EP|UA"}, ...]
    Regole:
      - C: compresenza con assente (stessa classe/ora), esclusi assenti/usati/DISP
      - D: docenti con DISP a orario (lezioni.subject == 'DISP') in quell'ora, esclusi assenti/usati
      - EP/UA: docenti fittizi, sempre presenti, non soggetti a "già usato"
    """
    candidati = []

    # docenti già usati in quell'ora
    usati_ids = {x[0] for x in db.session.query(Sostituzione.docente_sostituto_id)
                 .filter(Sostituzione.data == giorno_data, Sostituzione.hour == hour_tag)
                 .all()}

    # docenti fittizi (EP/UA)
    speciali = (Docente.query
                .filter(Docente.attivo == True)
                .filter(getattr(Docente, "fittizio") == 1)
                .order_by(Docente.cognome, Docente.nome)
                .all())

    # -------- C: compresenza con l'assente --------
    if classe:
        lez_in_classe = (Lezione.query
                         .filter(Lezione.day == giorno_cod,
                                 Lezione.hour == hour_tag,
                                 Lezione.students_set == classe)
                         .all())

        visti = set()
        for lez in lez_in_classe:
            did = lez.docente_id

            if did == docente_assente_id:
                continue
            if did in assenti_ids:
                continue
            if did in usati_ids:
                continue

            # evita di considerare "DISP" come compresenza
            if lez.subject == "DISP":
                continue

            doc = Docente.query.get(did)
            if not doc or not doc.attivo:
                continue
            if bool(getattr(doc, "fittizio", False)):
                continue

            if did not in visti:
                visti.add(did)
                candidati.append({"docente": doc, "tipo": "C"})

    # -------- D: DISP a orario (subject == 'DISP') --------
    lez_disp = (Lezione.query
                .filter(Lezione.day == giorno_cod,
                        Lezione.hour == hour_tag,
                        Lezione.subject == "DISP")
                .all())

    visti = set()
    for lez in lez_disp:
        did = lez.docente_id
        if did in assenti_ids:
            continue
        if did in usati_ids:
            continue

        doc = Docente.query.get(did)
        if not doc or not doc.attivo:
            continue
        if bool(getattr(doc, "fittizio", False)):
            continue

        if did not in visti:
            visti.add(did)
            candidati.append({"docente": doc, "tipo": "D"})

    # -------- EP/UA: sempre disponibili --------
    for sp in speciali:
        # tipo stampato: "EP" o "UA" (uso il nome breve se lo hai messo così)
        # Se preferisci, puoi forzarlo con: sp.nome in ("EP","UA")
        tipo_sp = (sp.nome or "").strip().upper() or "SP"
        candidati.append({"docente": sp, "tipo": tipo_sp})

    # ordinamento semplice: C, D, EP, UA, poi cognome
    ordine = {"C": 0, "D": 1, "EP": 2, "UA": 3}
    candidati.sort(key=lambda x: (
        ordine.get(x["tipo"], 9),
        (x["docente"].cognome or ""),
        (x["docente"].nome or "")
    ))

    return candidati


def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'utente_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return wrapper


def role_required(*roles):
    """
    Permette accesso solo se session['ruolo'] è uno dei ruoli ammessi.
    """
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            if session.get('ruolo') not in roles:
                flash("Accesso non autorizzato.", "warning")
                return redirect(url_for('index'))
            return f(*args, **kwargs)
        return wrapper
    return decorator


def codice_giorno_da_data(data_ass: date) -> str:
    """
    Converte una data Python in codice giorno dell'orario (LUN, MAR, MER, GIO, VEN, SAB, DOM).
    weekday(): 0 = lunedì, 6 = domenica.
    """
    mappa = {
        0: 'LUN',
        1: 'MAR',
        2: 'MER',
        3: 'GIO',
        4: 'VEN',
        5: 'SAB',
        6: 'DOM',
    }
    return mappa[data_ass.weekday()]


def ore_str_to_list(ore_str: str):
    """
    '1,2,3' -> ['1','2','3'] (pulite)
    """
    out = []
    for x in (ore_str or "").split(","):
        x = x.strip()
        if x:
            out.append(x)
    return out


def ora_num_to_hour_tag(n: str) -> str:
    """
    '5' -> 'H5'
    """
    n = str(n).strip()
    return f"H{n}"


# ==================
# ROTTE AUTENTICAZIONE
# ==================

@app.route('/init_admin')
def init_admin():
    """
    Rotta una tantum per creare l'utente admin.
    Dopo averla usata con successo, ELIMINA o COMMENTA questa funzione.
    """
    username = 'admin'
    password = 'admin123'  # cambiala subito dopo il primo login

    if Utente.query.filter_by(username=username).first():
        return "Admin esiste già."

    hash_pw = generate_password_hash(password)
    admin = Utente(username=username, password_hash=hash_pw, ruolo='admin')
    db.session.add(admin)
    db.session.commit()
    return "Utente admin creato (username: admin, password: admin123). Ricorda di cambiare la password."


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        utente = Utente.query.filter_by(username=username).first()
        if utente and utente.verifica_password(password):
            session.clear()
            session['utente_id'] = utente.id
            session['username'] = utente.username
            session['ruolo'] = utente.ruolo
            # ✅ dopo login vai alla HOME con i 3 bottoni
            return redirect(url_for('index'))
        else:
            flash('Credenziali errate', 'danger')

    return render_template('login.html')


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# ==================
# ROTTE PRINCIPALI
# ==================

@app.route('/')
@login_required
def index():
    """
    Home: mostra i 3 bottoni (logica di abilitazione la fai in index.html)
    """
    return render_template('index.html')


# ==================
# GESTIONE ASSENZE
# ==================

from datetime import date, timedelta

@app.route('/assenze', methods=['GET', 'POST'])
@login_required
@role_required('segreteria', 'sostituzioni', 'admin')
def gestione_assenze():
    docenti = (Docente.query
               .filter_by(attivo=True)
               .order_by(Docente.cognome, Docente.nome)
               .all())

    # -----------------
    # INSERIMENTO (POST)
    # -----------------
    if request.method == 'POST':
        try:
            docente_id = (request.form.get('docente_id') or "").strip()
            dal_str = (request.form.get('data_inizio') or "").strip()   # <-- OK
            al_str  = (request.form.get('data_fine') or "").strip()     # <-- OK
            note_val = (request.form.get('note') or "").strip()

            if not docente_id or not dal_str:
                flash("Seleziona docente e data.", "warning")
                return redirect(url_for('gestione_assenze'))

            docente_id_int = int(docente_id)

            y, m, d = map(int, dal_str.split('-'))
            dal_date = date(y, m, d)

            if al_str:
                y2, m2, d2 = map(int, al_str.split('-'))
                al_date = date(y2, m2, d2)
            else:
                al_date = dal_date

            if al_date < dal_date:
                dal_date, al_date = al_date, dal_date

            aggiornite = 0
            inserite = 0
            senza_orario = []

            giorno_corr = dal_date
            while giorno_corr <= al_date:
                codice_g = codice_giorno_da_data(giorno_corr)

                lezioni = (Lezione.query
                           .filter_by(docente_id=docente_id_int, day=codice_g)
                           .all())

                ore_orario_set = set()
                for l in lezioni:
                    digits = ''.join(ch for ch in (l.hour or '') if ch.isdigit())
                    if digits:
                        ore_orario_set.add(str(int(digits)))  # normalizza

                if not ore_orario_set:
                    senza_orario.append(giorno_corr.strftime('%d/%m/%Y'))
                    giorno_corr += timedelta(days=1)
                    continue

                esistenti = (Assenza.query
                             .filter_by(docente_id=docente_id_int, data=giorno_corr)
                             .order_by(Assenza.id.asc())
                             .all())

                ore_esistenti = set()
                for a in esistenti:
                    for o in (a.ore or "").split(','):
                        o = o.strip()
                        if o:
                            try:
                                ore_esistenti.add(str(int(o)))
                            except ValueError:
                                pass

                ore_combinate = sorted(ore_esistenti.union(ore_orario_set), key=lambda x: int(x))
                ore_str = ",".join(ore_combinate)

                if esistenti:
                    principale = esistenti[0]
                    principale.ore = ore_str

                    if note_val:
                        principale.note = (principale.note + " | " + note_val).strip() if principale.note else note_val

                    for extra in esistenti[1:]:
                        db.session.delete(extra)

                    aggiornite += 1
                else:
                    nuova = Assenza(
                        docente_id=docente_id_int,
                        data=giorno_corr,
                        ore=ore_str,
                        note=note_val
                    )
                    db.session.add(nuova)
                    inserite += 1

                giorno_corr += timedelta(days=1)

            db.session.commit()

            if inserite or aggiornite:
                flash(f"Assenze salvate. Inserite: {inserite}, aggiornate: {aggiornite}.", "success")

            if senza_orario:
                flash("Nessuna ora a orario per: " + ", ".join(senza_orario), "warning")

            return redirect(url_for('gestione_assenze'))

        except Exception as e:
            db.session.rollback()
            flash(f"Errore inserimento assenza: {e}", "danger")
            return redirect(url_for('gestione_assenze'))

    # -----------------
    # FILTRI (GET)
    # -----------------
    f_docente_id = (request.args.get('f_docente_id') or "").strip()
    f_data = (request.args.get('f_data') or "").strip()
    view_mode = (request.args.get('view') or "compact").strip()

    query = Assenza.query

    if f_docente_id:
        try:
            query = query.filter(Assenza.docente_id == int(f_docente_id))
        except ValueError:
            pass

    if f_data:
        try:
            y, m, d = map(int, f_data.split('-'))
            query = query.filter(Assenza.data == date(y, m, d))
        except ValueError:
            pass

    assenze = query.order_by(Assenza.data.desc(), Assenza.id.desc()).all()

    return render_template(
        'assenze.html',
        docenti=docenti,
        assenze=assenze,
        f_docente_id=f_docente_id,
        f_data=f_data,
        view_mode=view_mode,
        data_oggi=date.today().isoformat()
    )


# 🔹 Cancella una singola ora da un'assenza
@app.route('/assenze/cancella', methods=['POST'])
@login_required
@role_required('segreteria', 'sostituzioni', 'admin')
def cancella_assenza_ora():
    assenza_id = request.form.get('assenza_id')
    ora = request.form.get('ora')

    if not assenza_id or not ora:
        flash("Dati mancanti per la cancellazione.", 'warning')
        return redirect(url_for('gestione_assenze'))

    assenza = Assenza.query.get(assenza_id)
    if not assenza:
        flash("Assenza non trovata.", 'danger')
        return redirect(url_for('gestione_assenze'))

    ore_list = [x for x in (assenza.ore or '').split(',') if x]

    if ora in ore_list:
        ore_list.remove(ora)

    if ore_list:
        assenza.ore = ",".join(ore_list)
        db.session.commit()
        flash(f"Ora {ora} rimossa dall'assenza.", 'success')
    else:
        db.session.delete(assenza)
        db.session.commit()
        flash("Assenza eliminata (nessuna ora residua).", 'success')

    return redirect(url_for('gestione_assenze'))


# ==================
# SOSTITUZIONI - DISPONIBILITÀ
# ==================
@app.route('/disponibilita')
@login_required
@role_required('sostituzioni', 'admin')
def disponibilita():
    data_str = request.args.get('data')
    if data_str:
        y, m, d = map(int, data_str.split('-'))
        giorno_data = date(y, m, d)
    else:
        giorno_data = date.today()
        data_str = giorno_data.isoformat()

    giorno_cod = codice_giorno_da_data(giorno_data)

    # assenze del giorno
    assenze_giorno = Assenza.query.filter_by(data=giorno_data).all()

    # assenti per ora: hour_tag -> set(docente_id)
    assenti_per_hour = {}
    for a in assenze_giorno:
        for ora_num in ore_str_to_list(a.ore):
            hour_tag = ora_num_to_hour_tag(ora_num)
            assenti_per_hour.setdefault(hour_tag, set()).add(a.docente_id)

    # sostituzioni già inserite nel giorno (per mostrare assegnato)
    sost_giorno = (Sostituzione.query
                   .filter(Sostituzione.data == giorno_data)
                   .all())

    sost_map = {}
    for sst in sost_giorno:
        key = (sst.hour, (sst.classe or ""), sst.docente_assente_id)
        sost_map[key] = sst

    # costruzione scoperture
    by_hour = {}  # hour_tag -> list of scoperture

    for a in assenze_giorno:
        doc_ass = Docente.query.get(a.docente_id)
        nome_ass = f"{doc_ass.cognome} {doc_ass.nome}".strip() if doc_ass else f"ID {a.docente_id}"

        for ora_num in ore_str_to_list(a.ore):
            hour_tag = ora_num_to_hour_tag(ora_num)

            lez = Lezione.query.filter_by(
                docente_id=a.docente_id,
                day=giorno_cod,
                hour=hour_tag
            ).first()

            # ✅ se il docente assente in quell'ora era DISP, NON è una scopertura
            if lez and lez.subject == "DISP":
                continue

            classe = (lez.students_set if lez else "") or ""

            scopertura = {
                "hour": hour_tag,
                "docente_assente_id": a.docente_id,
                "docente_assente_nome": nome_ass,
                "classe": classe,
                "candidati": [],
                "dropdown_compresenza": [],
                "dropdown_altri": [],
                "sostituzione": None
            }

            by_hour.setdefault(hour_tag, []).append(scopertura)

    # calcolo candidati + dropdown + sostituzione assegnata
    for hour_tag, scoperture in by_hour.items():
        assenti_ids = assenti_per_hour.get(hour_tag, set())

        # docenti già usati in quell'ora (per dropdown: escludiamo i reali già usati)
        usati_ids = {x[0] for x in db.session.query(Sostituzione.docente_sostituto_id)
                     .filter(Sostituzione.data == giorno_data, Sostituzione.hour == hour_tag)
                     .all()}

        # tutti docenti attivi
        tutti_attivi = (Docente.query
                        .filter(Docente.attivo == True)
                        .order_by(Docente.cognome, Docente.nome)
                        .all())

        for s in scoperture:
            # sostituzione già presente?
            key = (hour_tag, (s["classe"] or ""), s["docente_assente_id"])
            sst = sost_map.get(key)
            if sst:
                doc_sost = Docente.query.get(sst.docente_sostituto_id)
                sost_nome = f"{doc_sost.cognome} {doc_sost.nome}".strip() if doc_sost else f"ID {sst.docente_sostituto_id}"
                s["sostituzione"] = {
                    "id": sst.id,
                    "tipo": sst.tipo,
                    "sostituto_nome": sost_nome
                }

            # candidati automatici (C/D + EP/UA sempre)
            s["candidati"] = candidati_per_scopertura(
                giorno_data=giorno_data,
                giorno_cod=giorno_cod,
                hour_tag=hour_tag,
                classe=s["classe"],
                docente_assente_id=s["docente_assente_id"],
                assenti_ids=assenti_ids
            )

            # dropdown (manuale): compresenza prima, poi altri
            comp_ids = set()
            comp_docenti = []

            if s["classe"]:
                lez_comp = (Lezione.query
                            .filter(Lezione.day == giorno_cod,
                                    Lezione.hour == hour_tag,
                                    Lezione.students_set == s["classe"])
                            .all())

                for lez in lez_comp:
                    did = lez.docente_id
                    if did == s["docente_assente_id"]:
                        continue
                    if did in assenti_ids:
                        continue
                    if did in usati_ids:
                        continue
                    if lez.subject == "DISP":
                        continue

                    doc = Docente.query.get(did)
                    if doc and doc.attivo and not bool(getattr(doc, "fittizio", False)) and doc.id not in comp_ids:
                        comp_ids.add(doc.id)
                        comp_docenti.append(doc)

            comp_docenti.sort(key=lambda x: (x.cognome or "", x.nome or ""))

            # altri: tutti attivi, esclusi assenti; esclusi reali già usati; includi fittizi sempre
            altri = []
            for d in tutti_attivi:
                if d.id in assenti_ids:
                    continue
                if not bool(getattr(d, "fittizio", False)) and d.id in usati_ids:
                    continue
                if d.id in comp_ids:
                    continue
                altri.append(d)

            s["dropdown_compresenza"] = comp_docenti
            s["dropdown_altri"] = altri
    tutti_docenti = (Docente.query
                 .order_by(Docente.cognome, Docente.nome)
                 .all())

    return render_template(
        'disponibilita.html',
        data=data_str,
        by_hour=by_hour,
        tutti_docenti=tutti_docenti
    )

  



@app.route('/sostituzioni/assegna', methods=['POST'])
@login_required
@role_required('sostituzioni', 'admin')
def assegna_sostituzione():
    try:
        # ---- input ----
        data_str = (request.form.get('data') or "").strip()
        hour = (request.form.get('hour') or "").strip()
        classe = (request.form.get('classe') or "").strip()
        docente_assente_id = (request.form.get('docente_assente_id') or "").strip()

        # tipo: D / C / EP / UA
        tipo = (request.form.get('tipo') or "D").strip().upper()

        # docente_sostituto_id può mancare (EP/UA)
        docente_sostituto_id_raw = (request.form.get('docente_sostituto_id') or "").strip()

        # ---- controlli minimi ----
        if not data_str or not hour or not docente_assente_id:
            flash("Dati mancanti (data/ora/docente assente).", "warning")
            return redirect(url_for('disponibilita', data=(data_str or date.today().isoformat())))

        # parse data
        y, m, d = map(int, data_str.split('-'))
        giorno_data = date(y, m, d)

        docente_assente_id = int(docente_assente_id)

        # EP/UA: NON richiedono docente_sostituto_id (lo salviamo NULL)
        is_epua = tipo in ("EP", "UA")

        # D/C: richiedono docente sostituto
        if not is_epua and not docente_sostituto_id_raw:
            flash("Seleziona un docente sostituto (oppure scegli EP/UA).", "warning")
            return redirect(url_for('disponibilita', data=giorno_data.isoformat()))

        docente_sostituto_id = None
        if (not is_epua) and docente_sostituto_id_raw:
            docente_sostituto_id = int(docente_sostituto_id_raw)

        # ---- blocco: non assegnare se l'ora dell'assente è DISP ----
        giorno_cod = codice_giorno_da_data(giorno_data)
        lez_assente = Lezione.query.filter_by(
            docente_id=docente_assente_id,
            day=giorno_cod,
            hour=hour
        ).first()

        if lez_assente and (lez_assente.subject == "DISP"):
            flash("Non puoi assegnare: il docente assente era in DISP in quell'ora.", "warning")
            return redirect(url_for('disponibilita', data=giorno_data.isoformat()))

        # ---- esiste già la sostituzione per quella specifica scopertura? (UPsert) ----
        existing = (Sostituzione.query
                    .filter_by(
                        data=giorno_data,
                        hour=hour,
                        classe=classe,
                        docente_assente_id=docente_assente_id
                    )
                    .first())

        # ---- vincolo: un docente reale NON può essere usato 2 volte nella stessa ora ----
        # EP/UA: sempre consentiti perché docente_sostituto_id = NULL
        if not is_epua:
            gia_usato = (Sostituzione.query
                         .filter_by(
                             data=giorno_data,
                             hour=hour,
                             docente_sostituto_id=docente_sostituto_id
                         )
                         .first())

            if gia_usato and (not existing or gia_usato.id != existing.id):
                flash("Errore: docente già usato in quell'ora.", "warning")
                return redirect(url_for('disponibilita', data=giorno_data.isoformat()))

        # ---- crea o aggiorna ----
        if existing:
            existing.docente_sostituto_id = docente_sostituto_id  # None per EP/UA
            existing.tipo = tipo
            db.session.commit()
            flash("Sostituzione aggiornata.", "success")
        else:
            nuova = Sostituzione(
                data=giorno_data,
                hour=hour,
                classe=classe,
                docente_assente_id=docente_assente_id,
                docente_sostituto_id=docente_sostituto_id,  # None per EP/UA
                tipo=tipo
            )
            db.session.add(nuova)
            db.session.commit()
            flash("Sostituzione inserita.", "success")

        return redirect(url_for('disponibilita', data=giorno_data.isoformat()))

    except Exception as e:
        db.session.rollback()
        flash(f"Errore salvataggio sostituzione: {e}", "danger")
        return redirect(url_for('disponibilita', data=(request.form.get('data') or date.today().isoformat())))




@app.route('/sostituzioni/<int:sost_id>/cancella', methods=['POST'])
@login_required
@role_required('sostituzioni', 'admin')
def cancella_sostituzione(sost_id):
    sost = Sostituzione.query.get_or_404(sost_id)
    data_str = request.form.get('data') or sost.data.isoformat()

    db.session.delete(sost)
    db.session.commit()

    flash("Sostituzione eliminata.", "success")
    return redirect(url_for('disponibilita', data=data_str))



@app.route('/lista_sostituzioni')
@login_required
@role_required('sostituzioni', 'admin')
def lista_sostituzioni():
    data_str = request.args.get('data')
    if data_str:
        y, m, d = map(int, data_str.split('-'))
        giorno_data = date(y, m, d)
    else:
        giorno_data = date.today()
        data_str = giorno_data.isoformat()

    sost = (Sostituzione.query
            .filter(Sostituzione.data == giorno_data)
            .order_by(Sostituzione.hour, Sostituzione.classe, Sostituzione.id)
            .all())

    # carico docenti una volta sola
    docenti = {d.id: d for d in Docente.query.all()}

    righe = []
    for s in sost:
        ass = docenti.get(s.docente_assente_id)

        tipo = (s.tipo or "").strip().upper()

        # Nome docente assente
        assente_nome = (
            f"{ass.cognome} {ass.nome}".strip()
            if ass else f"ID {s.docente_assente_id}"
        )

        # Nome sostituto:
        # - EP/UA: niente sostituto (docente_sostituto_id può essere NULL)
        # - D/C: mostra nome del sostituto se presente
        sostituto_nome = "—"
        if tipo not in ("EP", "UA"):
            if s.docente_sostituto_id is not None:
                sostit = docenti.get(s.docente_sostituto_id)
                sostituto_nome = (
                    f"{sostit.cognome} {sostit.nome}".strip()
                    if sostit else f"ID {s.docente_sostituto_id}"
                )

        righe.append({
            "id": s.id,
            "hour": s.hour,
            "classe": s.classe or "",
            "tipo": tipo,
            "assente_nome": assente_nome,
            "sostituto_nome": sostituto_nome,
            "entrata_posticipata": bool(getattr(s, "entrata_posticipata", False)),
            "uscita_anticipata": bool(getattr(s, "uscita_anticipata", False)),
        })

    return render_template(
        'lista_sostituzioni.html',
        data=data_str,
        righe=righe
    )



@app.route('/sostituzioni/<int:sost_id>/cancella_da_lista', methods=['POST'])
@login_required
@role_required('sostituzioni', 'admin')
def cancella_sostituzione_da_lista(sost_id):
    sost = Sostituzione.query.get_or_404(sost_id)
    data_str = request.form.get('data') or sost.data.isoformat()

    db.session.delete(sost)
    db.session.commit()

    flash("Sostituzione eliminata.", "success")
    return redirect(url_for('lista_sostituzioni', data=data_str))



from io import BytesIO
from datetime import date
from flask import request, send_file
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor




@app.route('/stampa_sostituzioni_word')
@login_required
@role_required('sostituzioni', 'admin')
def stampa_sostituzioni_word():
    data_str = request.args.get('data')

    if data_str:
        y, m, d = map(int, data_str.split('-'))
        giorno_data = date(y, m, d)
    else:
        giorno_data = date.today()
        data_str = giorno_data.isoformat()

    giorni = ["LUN", "MAR", "MER", "GIO", "VEN", "SAB", "DOM"]
    giorno_txt = giorni[giorno_data.weekday()]

    sost = (
        Sostituzione.query
        .filter(Sostituzione.data == giorno_data)
        .order_by(Sostituzione.hour, Sostituzione.classe, Sostituzione.id)
        .all()
    )

    docenti = {d.id: d for d in Docente.query.all()}

    doc = Document()

    # ---------------- ORIENTAMENTO ORIZZONTALE ----------------
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # ---------------- FUNZIONI UTILI ----------------
    def set_cell_shading(cell, fill_hex: str):
        """fill_hex tipo 'FFF2CC' (senza #)"""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

    def set_row_height_1cm(row):
        row.height = Cm(1)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # ---------------- TITOLO ----------------
    titolo = doc.add_heading(
        f"SOSTITUZIONI – {giorno_txt} {giorno_data.strftime('%d/%m/%Y')}",
        level=1
    )
    titolo.alignment = 1

    # ---------------- TABELLA (NO COLONNA TIPO) ----------------
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    intestazioni = ["ORA", "CLASSE", "DOCENTE ASSENTE", "SOSTITUTO", "FIRMA SOSTITUTO"]

    for i, txt in enumerate(intestazioni):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], "D9E1F2")  # azzurro pastello

    # intestazione ripetuta su ogni pagina
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))
    set_row_height_1cm(table.rows[0])

    # ---------------- RIGHE ----------------
    # colori alternati pastello (molto leggeri)
    ROW_A = "FFFFFF"
    ROW_B = "F2F2F2"  # grigio chiarissimo (pastello neutro)

    for idx, s in enumerate(sost):
        ass = docenti.get(s.docente_assente_id)
        sostit = docenti.get(s.docente_sostituto_id)

        assente_nome = f"{ass.cognome} {ass.nome}".strip() if ass else ""

        # flags EP/UA dal DB (in base a quanto mi hai detto)
        ep_flag = bool(getattr(s, "entrata_posticipata", False))
        ua_flag = bool(getattr(s, "uscita_anticipata", False))

        row = table.add_row()
        set_row_height_1cm(row)

        # shading alternato su tutta la riga (pastello)
        fill = ROW_A if (idx % 2 == 0) else ROW_B
        for c in row.cells:
            set_cell_shading(c, fill)

        c0, c1, c2, c3, c4 = row.cells

        # ORA
        c0.text = str(s.hour or "")

        # CLASSE
        c1.text = str(s.classe or "")

        # DOCENTE ASSENTE
        c2.text = assente_nome

        # SOSTITUTO (nome in neretto + sigla tra [] + EP/UA sotto)
        c3.text = ""
        p1 = c3.paragraphs[0]

        if sostit:
            nome = f"{sostit.cognome} {sostit.nome}".strip()

            # sigle accanto al nome (come avevi chiesto)
            sigla = ""
            if s.tipo == "D":
                sigla = "[D]"
            elif s.tipo == "C":
                sigla = "[C]"
            elif s.tipo == "*":
                sigla = "[*]"

            run_nome = p1.add_run(f"{nome} {sigla}".strip())
            run_nome.bold = True
            run_nome.font.size = Pt(10)
        else:
            # nessun sostituto: non scrivere "ID None"
            # lascia vuoto, ma sotto puoi comunque stampare EP/UA
            p1.add_run("—").font.size = Pt(10)

        # EP/UA sotto al sostituto (in piccolo)
        # (NOTA: li scrivo come richiesto, non come sigla)
        if ep_flag or ua_flag:
            p2 = c3.add_paragraph()
            p2_run = p2.add_run(
                "Entrata posticipata" if ep_flag else ""
            )
            if ep_flag and ua_flag:
                p2_run.text = "Entrata posticipata / Uscita anticipata"
            elif ua_flag:
                p2_run.text = "Uscita anticipata"

            p2_run.italic = True
            p2_run.font.size = Pt(9)

        # FIRMA
        c4.text = ""

    # ---------------- FONT GENERALE ----------------
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"sostituzioni_{giorno_data.strftime('%Y-%m-%d')}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Rotta di debug opzionale per vedere l'orario di un docente
@app.route('/debug_orario/<int:docente_id>')
@login_required
def debug_orario(docente_id):
    lez = Lezione.query.filter_by(docente_id=docente_id).order_by(Lezione.day, Lezione.hour).all()
    out = []
    for l in lez:
        out.append(f"{l.day} {l.hour} {l.students_set} {l.subject} ({l.room})")
    return "<br>".join(out) or "Nessuna lezione"


if __name__ == '__main__':
    app.run(debug=True)
